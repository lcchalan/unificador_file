# lector_file.py
# -*- coding: utf-8 -*-
"""
Lógica de procesamiento para combinar contenidos de varios .docx por títulos (H1/H2/H3),
pensada para ejecutarse en servidor (stateless, todo en memoria).

APIs expuestas:
- headings_from_docx(content_bytes) -> List[Dict]
- procesar(archivos, niveles, titulos_whitelist, enforce_whitelist=False) -> Dict[str, bytes]
- procesar_grouped(archivos, level, titulos_whitelist, enforce_whitelist=False) -> Dict[str, bytes]

Donde:
- archivos: List[{"name": str, "content": bytes}]
- niveles: List[int]   (ej. [1,2,3])
- titulos_whitelist: List[str] (si vacío => no se filtra por texto)
- enforce_whitelist: si True, sólo incluye títulos exactamente en whitelist; si False, whitelist sirve de filtro opcional.

Dependencias: python-docx, openpyxl, lxml
"""

from __future__ import annotations
import io
import re
import unicodedata
from typing import List, Dict, Tuple, Optional, Iterable

from docx import Document
from docx.document import Document as _DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

import pandas as pd

# ---------------------------
# Utilidades
# ---------------------------

def _to_docx(content_bytes: bytes) -> _DocxDocument:
    bio = io.BytesIO(content_bytes)
    return Document(bio)

def _save_docx_to_bytes(doc: _DocxDocument) -> bytes:
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def _normalize_text(s: str) -> str:
    if s is None:
        return ""
    s2 = unicodedata.normalize("NFKC", s)
    return re.sub(r"\s+", " ", s2).strip()

def _is_heading(paragraph: Paragraph) -> Tuple[bool, Optional[int]]:
    """
    Detecta si un párrafo es encabezado y devuelve (es_heading, nivel).
    Soporta estilos en ES/EN típicos (Título 1/Heading 1, etc).
    """
    style = paragraph.style.name if paragraph.style is not None else ""
    txt = (paragraph.text or "").strip()

    # Preferencia: documenta el nivel directo si lo provee python-docx
    # (No siempre disponible, depende de la plantilla)
    # Fallback: estilo por nombre.
    level = None

    # Reglas por nombre de estilo (comunes en Word ES/EN)
    style_norm = style.lower()
    m = re.search(r"(heading|encabezado|t[íi]tulo)\s*(\d+)", style_norm)
    if m:
        try:
            level = int(m.group(2))
        except Exception:
            level = None

    # A veces hay estilos "Heading 1 Char" (caracter), ignorar
    if "char" in style_norm and level is not None:
        # si es Char, no lo consideramos heading estructural
        return False, None

    if level is not None and 1 <= level <= 9:
        return True, level

    # Último recurso: si el párrafo tiene outline_level en el XML (no siempre accesible)
    # (python-docx no lo expone fácilmente; lo omitimos para robustez)

    return False, None

def _iter_block_items(parent) -> Iterable:
    """
    Itera en orden real del documento: párrafos y tablas.
    parent: Document o _Cell
    Retorna Paragraph o Table.
    """
    if isinstance(parent, _DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Tipo de padre no soportado para iteración")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _copy_paragraph(dst_doc: _DocxDocument, p: Paragraph):
    # Copiamos sólo texto y formato mínimo (negritas/cursivas subyacentes no se preservan completamente).
    # Para servidor estable, esto es suficiente.
    new_p = dst_doc.add_paragraph()
    text = p.text or ""
    if text:
        new_p.add_run(text)

def _copy_table_as_text(dst_doc: _DocxDocument, tbl: Table):
    """
    Como copiar tablas 1:1 en python-docx no es trivial (copiar XML),
    aquí volcamos cada fila como texto CSV-like. Suficiente para contexto.
    (Además exportamos las tablas reales a Excel por separado.)
    """
    rows = []
    for r in tbl.rows:
        row_vals = []
        for c in r.cells:
            row_vals.append(_normalize_text(c.text))
        rows.append(row_vals)
    if not rows:
        return
    # Insertamos un bloque con el "texto" de la tabla
    dst_doc.add_paragraph("")  # separador
    for row in rows:
        dst_doc.add_paragraph(" | ".join(row))
    dst_doc.add_paragraph("")  # separador

def _sanitize_filename(name: str) -> str:
    name = _normalize_text(name)
    name = re.sub(r'[\\/:*?"<>|]+', "_", name)
    name = name.strip(" .")
    if not name:
        name = "archivo"
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name

# ---------------------------
# Extracción de encabezados
# ---------------------------

def headings_from_docx(content_bytes: bytes) -> List[Dict]:
    """
    Retorna lista de dicts: [{"level": int, "text": str}, ...] en orden.
    """
    doc = _to_docx(content_bytes)
    out = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            is_h, lvl = _is_heading(block)
            if is_h:
                out.append({"level": int(lvl), "text": _normalize_text(block.text)})
    return out

# ---------------------------
# Segmentación por secciones de encabezado
# ---------------------------

def _split_sections_by_levels(doc: _DocxDocument, include_levels: List[int]) -> List[Dict]:
    """
    Divide el documento en secciones a partir de encabezados cuyos niveles estén en include_levels.
    Cada sección = {"level": int, "title": str, "content": List[Union[Paragraph, Table]]}
    El contenido incluye los bloques desde el título hasta ANTES del siguiente título de nivel
    igual o menor (clásico "sección").
    """
    sections: List[Dict] = []
    current = None
    current_level_stack: List[int] = []

    # Recopilamos todos los encabezados (para cortes)
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            is_h, lvl = _is_heading(block)
            if is_h:
                title = _normalize_text(block.text)

                # Si el nivel del heading está en los que consideramos "inicio de sección"
                # abrimos una nueva sección
                if lvl in include_levels:
                    # cerrar la actual
                    if current:
                        sections.append(current)
                    current = {"level": int(lvl), "title": title, "content": []}
                    current_level_stack = [lvl]
                    continue

                # Si es encabezado pero no está en include_levels, podría marcar fin de la sección
                # Si aparece un heading de nivel <= nivel de la sección actual, cerramos
                if current and lvl <= current["level"]:
                    sections.append(current)
                    current = None
                    current_level_stack = []

                # Si no estamos dentro de sección que capture, seguimos
                # (los contenidos fuera de secciones seleccionadas no se incluyen)
            else:
                # párrafo normal
                if current:
                    current["content"].append(block)

        elif isinstance(block, Table):
            if current:
                current["content"].append(block)

    if current:
        sections.append(current)

    return sections

def _collect_tables_from_sections(sections: List[Dict]) -> List[Tuple[str, List[List[str]]]]:
    """
    Extrae tablas de las secciones (como matrices de texto).
    Retorna lista de tuplas: (title, rows[list[list[str]]])
    """
    out = []
    for sec in sections:
        title = sec["title"]
        for b in sec["content"]:
            if isinstance(b, Table):
                rows = []
                for r in b.rows:
                    row_vals = []
                    for c in r.cells:
                        row_vals.append(_normalize_text(c.text))
                    rows.append(row_vals)
                if rows:
                    out.append((title, rows))
    return out

# ---------------------------
# Export: Excel (tablas)
# ---------------------------

def _excel_from_tables(all_tables: List[Tuple[str, List[List[str]]]]) -> bytes:
    """
    Crea un Excel con una hoja por tabla encontrada.
    Hoja: "Tabla_001 - <titulo recortado>"
    """
    if not all_tables:
        # Excel vacío con hoja "Tablas" indicando que no hay
        with pd.ExcelWriter(io.BytesIO(), engine="openpyxl") as writer:
            df = pd.DataFrame([{"info": "No se encontraron tablas"}])
            df.to_excel(writer, index=False, sheet_name="Tablas")
            bio = writer.book.properties  # fuerza escritura
        # Re-crear para obtener bytes (truco)
        bio2 = io.BytesIO()
        with pd.ExcelWriter(bio2, engine="openpyxl") as writer2:
            df.to_excel(writer2, index=False, sheet_name="Tablas")
        return bio2.getvalue()

    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        for idx, (title, rows) in enumerate(all_tables, start=1):
            # Normalizamos en DataFrame (relleno desigual)
            max_cols = max(len(r) for r in rows) if rows else 0
            norm_rows = [r + [""]*(max_cols - len(r)) for r in rows]
            df = pd.DataFrame(norm_rows)
            # nombre de hoja
            t = _normalize_text(title)
            t = (t[:22] + "…") if len(t) > 23 else t
            sheet = f"Tabla_{idx:03d}"
            if t:
                sheet = f"{sheet}_{t}"
            # Excel limita a 31 chars
            sheet = sheet[:31]
            df.to_excel(writer, index=False, header=False, sheet_name=sheet)
    return bio.getvalue()

# ---------------------------
# Export: DOCX (unificado / agrupado)
# ---------------------------

def _append_section_to_doc(dst: _DocxDocument, source_name: str, section: Dict):
    # Encabezado de documento → título de la sección
    # Añadimos también el nombre del archivo fuente (pequeño) para trazabilidad
    title = section["title"]
    lvl = section["level"]

    # Heading apropiado (Word usa 0-9; python-docx admite 0..9)
    # Usamos nivel 1 siempre para visibilidad, y ponemos (Hn) en texto
    dst.add_heading(f"{title}", level=1)
    dst.add_paragraph(f"[Fuente: {source_name}]").italic = True

    for b in section["content"]:
        if isinstance(b, Paragraph):
            _copy_paragraph(dst, b)
        elif isinstance(b, Table):
            _copy_table_as_text(dst, b)
    dst.add_paragraph("")  # separador

def _merge_unificado(archivos: List[Dict], niveles: List[int],
                     titles_whitelist: List[str], enforce_whitelist: bool) -> Tuple[bytes, bytes]:
    """
    Crea un único DOCX con todas las secciones coincidentes de todos los documentos.
    Además retorna un Excel con todas las tablas encontradas en esas secciones.
    """
    out_doc = Document()
    out_doc.add_heading("Documento Unificado", level=0)

    all_tables: List[Tuple[str, List[List[str]]]] = []

    wl_norm = set(_normalize_text(t) for t in (titles_whitelist or []))

    for item in archivos:
        src_name = item.get("name", "archivo.docx")
        doc = _to_docx(item["content"])
        sections = _split_sections_by_levels(doc, niveles)

        # Filtrado por whitelist (si corresponde)
        if enforce_whitelist and wl_norm:
            sections = [s for s in sections if _normalize_text(s["title"]) in wl_norm]
        elif wl_norm:
            # Modo filtro-suave: si hay whitelist (no vacía), incluimos sólo los títulos listados
            sections = [s for s in sections if _normalize_text(s["title"]) in wl_norm]

        # Volcado
        for sec in sections:
            _append_section_to_doc(out_doc, src_name, sec)

        # Tablas para Excel
        all_tables.extend(_collect_tables_from_sections(sections))

    docx_bytes = _save_docx_to_bytes(out_doc)
    xlsx_bytes = _excel_from_tables(all_tables)
    return docx_bytes, xlsx_bytes

def _merge_grouped_by_title(archivos: List[Dict], level: int,
                            titles_whitelist: List[str], enforce_whitelist: bool) -> Dict[str, bytes]:
    """
    Crea un DOCX por cada título (del nivel dado) agregando el contenido proveniente
    de todos los documentos que tengan ese título.
    Retorna: {"<titulo>.docx": bytes, ...}
    """
    wl_norm = set(_normalize_text(t) for t in (titles_whitelist or []))
    # Mapa título -> lista de secciones (de varios docs)
    grouped: Dict[str, List[Tuple[str, Dict]]] = {}

    for item in archivos:
        src_name = item.get("name", "archivo.docx")
        doc = _to_docx(item["content"])
        sections = _split_sections_by_levels(doc, [level])

        for sec in sections:
            tnorm = _normalize_text(sec["title"])
            if enforce_whitelist and wl_norm and tnorm not in wl_norm:
                continue
            if wl_norm and tnorm not in wl_norm:
                # filtro-suave: si se proporcionó whitelist, incluimos solo listados
                continue
            grouped.setdefault(tnorm, []).append((src_name, sec))

    # Construcción de documentos finales
    out: Dict[str, bytes] = {}
    for title, items in grouped.items():
        if not title:
            continue
        d = Document()
        d.add_heading(title, level=0)
        for (src_name, sec) in items:
            _append_section_to_doc(d, src_name, sec)
        out[_sanitize_filename(f"{title}.docx")] = _save_docx_to_bytes(d)

    return out

# ---------------------------
# Funciones públicas
# ---------------------------

def procesar(archivos: List[Dict],
             niveles: List[int],
             titulos_whitelist: List[str],
             enforce_whitelist: bool = False) -> Dict[str, bytes]:
    """
    Modo UNIFICADO: regresa {"unificado.docx": bytes, "tablas.xlsx": bytes}
    - niveles: ej. [1,2,3]
    - titulos_whitelist: [] para todos; si no vacío, se usa como filtro de inclusión.
    - enforce_whitelist: si True, sólo incluye exactamente títulos en whitelist.
    """
    docx_bytes, xlsx_bytes = _merge_unificado(archivos, niveles, titulos_whitelist, enforce_whitelist)
    return {
        "unificado.docx": docx_bytes,
        "tablas.xlsx": xlsx_bytes
    }

def procesar_grouped(archivos: List[Dict],
                     level: int,
                     titulos_whitelist: List[str],
                     enforce_whitelist: bool = False) -> Dict[str, bytes]:
    """
    Modo POR TÍTULO: regresa {"<titulo>.docx": bytes, ...} agrupando por ese nivel.
    """
    return _merge_grouped_by_title(archivos, level, titulos_whitelist, enforce_whitelist)
def _merge_grouped_by_title(archivos: List[Dict], title_exact: str) -> Dict[str, bytes]:
    """
    Crea 1 DOCX con el contenido de TODOS los documentos que tengan ese título (en cualquier nivel).
    Retorna {"<titulo>.docx": bytes}
    """
    tnorm = _normalize_text(title_exact)
    d = Document()
    d.add_heading(title_exact, level=0)

    got = False
    for item in archivos:
        src = item.get("name", "archivo.docx")
        doc = _to_docx(item["content"])
        # buscamos secciones para TODOS los niveles (1..9)
        sections = _split_sections_by_levels(doc, list(range(1,10)))
        for sec in sections:
            if _normalize_text(sec["title"]) == tnorm:
                _append_section_to_doc(d, src, sec)
                got = True

    if not got:
        return {}
    return { _sanitize_filename(f"{title_exact}.docx"): _save_docx_to_bytes(d) }

# ---------------------------
# Modo CLI local (opcional)
# ---------------------------

def _demo_main():
    """
    Ejecución local de prueba (opcional). No se llama en servidor.
    - Lee .docx de ./data_in
    - Genera ./_out/unificado.docx y ./_out/tablas.xlsx
      y ./_out/por_titulo/<titulo>.docx
    """
    import os
    from glob import glob
    in_dir = "./data_in"
    out_dir = "./_out"
    os.makedirs(out_dir, exist_ok=True)
    files = []
    for p in glob(os.path.join(in_dir, "*.docx")):
        with open(p, "rb") as f:
            files.append({"name": os.path.basename(p), "content": f.read()})

    if not files:
        print("No hay .docx en ./data_in")
        return

    # Unificado
    uni = procesar(files, [1,2,3], [], enforce_whitelist=False)
    with open(os.path.join(out_dir, "unificado.docx"), "wb") as f:
        f.write(uni["unificado.docx"])
    with open(os.path.join(out_dir, "tablas.xlsx"), "wb") as f:
        f.write(uni["tablas.xlsx"])
    print("Generado unificado.docx y tablas.xlsx")

    # Por título (H1)
    grouped = procesar_grouped(files, 1, [], enforce_whitelist=False)
    folder = os.path.join(out_dir, "por_titulo")
    os.makedirs(folder, exist_ok=True)
    for fn, data in grouped.items():
        with open(os.path.join(folder, fn), "wb") as f:
            f.write(data)
    print(f"Generados {len(grouped)} archivos por título (H1).")

if __name__ == "__main__":
    _demo_main()
