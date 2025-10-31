import os
import re
import shutil
import tempfile
import unicodedata
from collections import defaultdict
from typing import List, Dict, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

# ----------------------
# Rutas
# ----------------------
ruta_entrada = "/Users/utpl/Desktop/doc"   # DOCX de entrada
ruta_salida  = "/Users/utpl/Desktop/doc1"  # salida DOCX/XLSX + reportes

os.makedirs(ruta_salida, exist_ok=True)

# ----------------------
# Índice de TÍTULOS (con numeración)
# ----------------------
TITULOS_RAW = """
01. Plan de formación integral del estudiante
06. Plan de admisión, acogida y acompañamiento académico de estudiantes
23. Plan de seguimiento y mejora de indicadores del perfil docente
25. Plan de formación integral del docente
26. Plan de mejora del proceso de evaluación integral docente 
03. Plan implantación del marco de competencias UTPL
04. Plan de prospectiva y creación de nueva oferta
07. Plan de acciones curriculares para el fortalecimiento de las competencias genéricas
11. Plan de fortalecimiento de prácticas preprofesionales y proyectos de vinculación
12. Plan de fortalecimiento de criterios para la evaluación de la calidad de carreras y programas académicos
13. Plan de acciones curriculares para el fortalecimiento de la empleabilidad del graduado UTPL
16. Plan de mejora del proceso de elaboración y seguimiento de planes docentes
18. Plan de mejora de ambientes de aprendizaje
19. Plan de mejora de evaluación de los aprendizajes
20. Plan de mejora del proceso de integración curricular
21. Plan de mejora del proceso de titulación
22. Plan de seguimiento y mejora de la labor tutorial
08. Plan de internacionalización del currículo
24. Plan de intervención de personal académico en territorio
05. Plan de acciones académicas orientadas a la comunicación y promoción de la oferta
09. Plan de innovación educativa
10. Plan de implantación de metodologías activas en el currículo
28. Plan de formación de líderes académicos 
29. Plan de posicionamiento institucional en innovación educativa
30. Plan de investigación sobre innovación educativa, EaD, MP
""".strip()

# ----------------------
# Utilidades
# ----------------------
def quitar_acentos(texto: str) -> str:
    if texto is None:
        return ""
    return "".join(c for c in unicodedata.normalize("NFD", texto) if unicodedata.category(c) != "Mn")

def norm_str(s: str) -> str:
    s = (s or "").strip()
    s = quitar_acentos(s.lower())
    s = re.sub(r"\s+", " ", s)
    return s

def quitar_numeracion_inicio(s: str):
    """Devuelve (numero_str|'' , texto_sin_numeracion). Acepta '01. ', '1) ', '12 - ', etc."""
    m = re.match(r"^\s*(\d{1,3})\s*[-\.\)]\s*(.*)$", s.strip())
    if m:
        return m.group(1), m.group(2).strip()
    return "", s.strip()

def es_docx_valido(nombre: str) -> bool:
    return nombre.lower().endswith(".docx") and not (nombre.startswith("~$") or nombre.startswith("."))

def iter_block_items(doc: Document):
    """Itera bloques del body (párrafos y tablas) preservando orden y devolviendo (tipo, bloque, elem)."""
    body = doc._element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith('p'):
            yield ('p', Paragraph(child, doc), child)
        elif child.tag.endswith('tbl'):
            yield ('t', Table(child, doc), child)

def extraer_tabla_2d(tbl: Table) -> List[List[str]]:
    data = []
    for row in tbl.rows:
        fila = []
        for cell in row.cells:
            fila.append((cell.text or "").strip())
        data.append(fila)
    return data

def sanitizar_nombre(nombre: str) -> str:
    nombre = (nombre or "").strip()
    nombre = re.sub(r"[\\/*?\"<>|:]", "_", nombre)
    nombre = re.sub(r"\s+", "_", nombre)
    return nombre[:180]

# ----------------------
# Preparar lista y mapa de títulos
# - Detectamos comparando texto SIN numeración (normalizado).
# - Mostramos y guardamos con numeración: "01 Plan de ..."
# ----------------------
titulos_display: List[str] = []               # "01 Plan de …" (con número, sin punto)
TITULOS_NORM_MAP: Dict[str, str] = {}         # norm("Plan de …") -> "01 Plan de …"

for linea in TITULOS_RAW.splitlines():
    linea = linea.strip()
    if not linea:
        continue
    num, texto = quitar_numeracion_inicio(linea)
    if not texto:
        continue
    display = (f"{int(num):02d} {texto}") if num else texto
    titulos_display.append(display)
    TITULOS_NORM_MAP[norm_str(texto)] = display

def es_titulo_de_indice(parrafo: Paragraph) -> str:
    """Si el párrafo coincide con un título (ignorando numeración/acento/caso), retorna el título DISPLAY (con número)."""
    txt = (parrafo.text or "").strip()
    if not txt:
        return ""
    _, sin_num = quitar_numeracion_inicio(txt)
    key = norm_str(sin_num)
    return TITULOS_NORM_MAP.get(key, "")

# ----------------------
# Paso 1: Indexar por TÍTULO (display) los índices de p/t por archivo y recolectar tablas
#         + matriz de presencia por archivo/título
# ----------------------
secciones_idx: Dict[str, Dict[str, List[int]]] = defaultdict(lambda: defaultdict(list))  # titulo_display -> archivo -> [idx_local_pt]
tablas_por_titulo: Dict[str, List[Tuple[str, List[List[str]]]]] = defaultdict(list)

archivos_analizados: List[str] = []
found_titles_by_file: Dict[str, set] = defaultdict(set)  # archivo -> {titulo_display encontrados}
archivos_sin_titulos: List[str] = []                    # archivos que no contenían ningún título

for archivo in sorted(os.listdir(ruta_entrada)):
    if not es_docx_valido(archivo):
        continue
    archivos_analizados.append(archivo)

    ruta_arch = os.path.join(ruta_entrada, archivo)
    try:
        doc = Document(ruta_arch)
    except Exception as e:
        print(f"⚠️ No se pudo abrir '{archivo}': {e}")
        continue

    titulo_actual = None
    any_title_found = False

    # lista local de SOLO p/t para indexar por posición local
    pt_elems: List[Tuple[str, object, object]] = []
    for tipo, bloque, elem in iter_block_items(doc):
        if tipo in ('p', 't'):
            pt_elems.append((tipo, bloque, elem))

    for local_idx, (tipo, bloque, elem) in enumerate(pt_elems):
        if tipo == 'p':
            maybe_title = es_titulo_de_indice(bloque)
            if maybe_title:
                titulo_actual = maybe_title
                any_title_found = True
                found_titles_by_file[archivo].add(maybe_title)  # <- MARCAMOS PRESENCIA
                _ = secciones_idx[titulo_actual][archivo]
                continue  # el párrafo del título no se guarda como contenido

        if titulo_actual:
            secciones_idx[titulo_actual][archivo].append(local_idx)
            if tipo == 't':
                tablas_por_titulo[titulo_actual].append((archivo, extraer_tabla_2d(bloque)))

    if not any_title_found:
        archivos_sin_titulos.append(archivo)

# ----------------------
# Paso 2: Por cada TÍTULO, construir DOCX uniendo fragmentos por archivo (docxcompose)
# ----------------------
for titulo_display, archivos_pos in secciones_idx.items():
    frag_paths: List[str] = []

    for archivo, posiciones_keep_local in archivos_pos.items():
        if not posiciones_keep_local:
            continue

        src_path = os.path.join(ruta_entrada, archivo)
        tmp_dir = tempfile.mkdtemp(prefix="frag_")
        frag_path = os.path.join(tmp_dir, f"frag_{archivo}")
        shutil.copyfile(src_path, frag_path)

        frag_doc = Document(frag_path)
        body = frag_doc._element.body

        # reconstruir lista SOLO de p/t en el fragmento
        pt_elems_doc: List[object] = []
        for ch in list(body.iterchildren()):
            if ch.tag.endswith('p') or ch.tag.endswith('tbl'):
                pt_elems_doc.append(ch)

        if not pt_elems_doc:
            continue

        keep_set = set(posiciones_keep_local)

        # borrar p/t no deseados (de atrás adelante)
        for local_idx in range(len(pt_elems_doc) - 1, -1, -1):
            if local_idx not in keep_set:
                ch = pt_elems_doc[local_idx]
                if ch in body:
                    body.remove(ch)

        # si quedó algo de contenido útil
        quedo_contenido = any(ch.tag.endswith('p') or ch.tag.endswith('tbl') for ch in body.iterchildren())
        if not quedo_contenido:
            continue

        # insertar rótulo "Nombre del archivo" al inicio
        p = frag_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(archivo)
        r.bold = True
        p_elem = p._p
        if p_elem in body:
            body.remove(p_elem)
        body.insert(0, p_elem)

        frag_doc.save(frag_path)
        frag_paths.append(frag_path)

    # unir fragmentos si hay
    out_docx = os.path.join(ruta_salida, f"{sanitizar_nombre(titulo_display)}.docx")
    if frag_paths:
        base = Document()
        composer = Composer(base)
        for fp in frag_paths:
            composer.append(Document(fp))
        composer.save(out_docx)
    else:
        d = Document()
        d.add_heading(titulo_display, level=1)
        d.add_paragraph("No se encontró contenido para este título en los archivos analizados.")
        d.save(out_docx)

    # Excel por título (solo tablas)
    wb = Workbook()
    ws = wb.active
    ws.title = "Tablas"
    fila = 1
    bold = Font(bold=True)
    center = Alignment(horizontal="center")
    total_tablas = 0

    if titulo_display in tablas_por_titulo:
        for archivo_tab, data in tablas_por_titulo[titulo_display]:
            if not data:
                continue
            ncols = max((len(r) for r in data), default=1)
            ws.merge_cells(start_row=fila, start_column=1, end_row=fila, end_column=max(1, ncols))
            c = ws.cell(row=fila, column=1, value=archivo_tab)
            c.font = bold
            c.alignment = center
            fila += 1
            for r in data:
                for j, val in enumerate(r, start=1):
                    ws.cell(row=fila, column=j, value=val)
                fila += 1
            fila += 1
            total_tablas += 1

    if total_tablas == 0:
        ws.cell(row=1, column=1, value="No se encontraron tablas en este título.")

    out_xlsx = os.path.join(ruta_salida, f"{sanitizar_nombre(titulo_display)}.xlsx")
    wb.save(out_xlsx)

# ----------------------
# Paso 3: Reporte detallado de ausencias
# ----------------------
reporte_path = os.path.join(ruta_salida, "reporte_ausencias_titulos.txt")
with open(reporte_path, "w", encoding="utf-8") as f:
    # A) Archivos sin ningún título
    f.write("A) Archivos sin ningún título del índice:\n")
    if archivos_sin_titulos:
        for nombre in archivos_sin_titulos:
            f.write(f"- {nombre}\n")
    else:
        f.write("  (Todos los archivos contienen al menos un título del índice)\n")
    f.write("\n")

    # B) Por archivo → títulos faltantes
    f.write("B) Por archivo → títulos faltantes:\n")
    all_titles_set = set(titulos_display)
    for archivo in archivos_analizados:
        found = found_titles_by_file.get(archivo, set())
        faltantes = sorted(all_titles_set - found)
        if faltantes:
            f.write(f"- {archivo}\n")
            for t in faltantes:
                f.write(f"    · Falta: {t}\n")
    f.write("\n")

    # C) Por título → archivos donde falta
    f.write("C) Por título → archivos donde falta:\n")
    files_set = set(archivos_analizados)
    for t in titulos_display:
        presentes = {file for file, titles in found_titles_by_file.items() if t in titles}
        faltan_en = sorted(files_set - presentes)
        if faltan_en:
            f.write(f"- {t}\n")
            for af in faltan_en:
                f.write(f"    · No está en: {af}\n")

print("✅ Proceso finalizado.")
print(f"📁 Reporte detallado: {reporte_path}")