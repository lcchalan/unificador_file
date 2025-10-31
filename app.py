# app.py
import os, io, uuid, shutil, zipfile
from typing import List, Dict
from flask import Flask, render_template, request, jsonify, send_file

# Lógica de Word/Excel (tu módulo servidor-stateless)
# Asegúrate de que lector_file.py esté en el mismo directorio
from lector_file import (
    procesar,                  # unificado  -> {"unificado.docx": bytes, "tablas.xlsx": bytes}
    headings_from_docx,        # extraer headings de un docx (bytes)
    _normalize_text,           # util: normaliza texto para comparación exacta
    _split_sections_by_levels, # util: dividir por secciones (todos los niveles)
    Document,                  # clase de python-docx para crear DOCX
    _save_docx_to_bytes,       # util: docx->bytes
    _sanitize_filename         # util: asegurar nombre archivo válido
)

app = Flask(__name__)

# Carpeta temporal para subidas (efímera en Render; persistente en local mientras corre)
TMP_ROOT = os.path.join(os.path.dirname(__file__), "tmp")
os.makedirs(TMP_ROOT, exist_ok=True)


# ---------------------------
# Helpers internos del web service
# ---------------------------
def _read_docx_from_folder(folder: str, include_subs: bool) -> List[Dict]:
    """
    Lee .docx desde el filesystem del servidor (solo útil si corres local o
    si tu servidor tiene esos archivos montados). Devuelve [{"name","content(bytes)"}].
    """
    files = []
    if not os.path.isdir(folder):
        return files

    def add_file(path, name):
        if name.lower().endswith(".docx") and not name.startswith("~$"):
            with open(path, "rb") as f:
                files.append({"name": name, "content": f.read()})

    if include_subs:
        for root, _, fnames in os.walk(folder):
            for fn in fnames:
                add_file(os.path.join(root, fn), fn)
    else:
        for fn in os.listdir(folder):
            p = os.path.join(folder, fn)
            if os.path.isfile(p):
                add_file(p, fn)
    return files


def _collect_overview_whitelist(archivos: List[Dict], whitelist: List[str]):
    """
    Calcula, por archivo, cuántos y cuáles títulos coinciden con la whitelist (texto exacto).
    Retorna lista de dicts: [{"name", "count", "titles":[...]}]
    """
    wl_norm = set(_normalize_text(t) for t in (whitelist or []))
    files_meta = []
    for a in archivos:
        name = a.get("name", "archivo.docx")
        matches = []
        try:
            hs_all = headings_from_docx(a["content"])
            for h in hs_all:
                tnorm = _normalize_text(h.get("text") or "")
                if tnorm in wl_norm:
                    matches.append(h.get("text") or "")
        except Exception:
            pass
        files_meta.append({"name": name, "count": len(matches), "titles": matches})
    files_meta.sort(key=lambda x: x["name"].casefold())
    return files_meta


def _merge_grouped_by_title_any_level(archivos: List[Dict], title_exact: str) -> Dict[str, bytes]:
    """
    Crea 1 DOCX que agrega el contenido de TODOS los documentos que tengan ese título,
    en cualquier nivel de heading. Devuelve {"<titulo>.docx": bytes}.
    """
    tnorm = _normalize_text(title_exact)
    d = Document()
    d.add_heading(title_exact, level=0)

    got = False
    for item in archivos:
        src = item.get("name", "archivo.docx")
        doc = Document(io.BytesIO(item["content"]))
        # Consideramos secciones para TODOS los niveles (1..9)
        sections = _split_sections_by_levels(doc, list(range(1, 10)))
        for sec in sections:
            if _normalize_text(sec["title"]) == tnorm:
                # Reutilizamos append de lector_file (misma lógica que unificado)
                # Copiamos contenido:
                # Encabezado
                d.add_heading(f"{sec['title']}", level=1)
                p = d.add_paragraph()
                run = p.add_run(f"[Fuente: {src}]")
                run.italic = True
                # Bloques
                for b in sec["content"]:
                    # imports locales para tipos (evitamos dependencias circulares aquí)
                    from docx.text.paragraph import Paragraph
                    from docx.table import Table
                    if isinstance(b, Paragraph):
                        newp = d.add_paragraph()
                        if b.text:
                            newp.add_run(b.text)
                    elif isinstance(b, Table):
                        # Volcado simple de tabla como texto legible
                        d.add_paragraph("")
                        for r in b.rows:
                            row_vals = []
                            for c in r.cells:
                                row_vals.append(_normalize_text(c.text))
                            d.add_paragraph(" | ".join(row_vals))
                        d.add_paragraph("")
                d.add_paragraph("")  # separador
                got = True

    if not got:
        return {}
    return {_sanitize_filename(f"{title_exact}.docx"): _save_docx_to_bytes(d)}


# ---------------------------
# Rutas
# ---------------------------
@app.get("/")
def home():
    return render_template("index.html")


@app.get("/health")
def health():
    return "OK", 200


@app.post("/api/scan-folder")
def api_scan_folder():
    """
    Escanea títulos en archivos .docx que existen en el servidor (ruta en 'folder'),
    y devuelve coincidencias contra 'whitelist'.
    Body JSON:
      {
        "folder": "/ruta/servidor",
        "include_subs": true/false,
        "whitelist": [ ...titulos... ]
      }
    """
    data = request.get_json(force=True)
    folder = (data.get("folder") or "").strip()
    include_subs = bool(data.get("include_subs", True))
    whitelist = data.get("whitelist") or []
    if not folder or not os.path.isdir(folder):
        return jsonify({"ok": False, "error": "Carpeta inválida."}), 400

    archivos = _read_docx_from_folder(folder, include_subs)
    files_meta = _collect_overview_whitelist(archivos, whitelist)
    return jsonify({
        "ok": True,
        "count": len(archivos),
        "files_meta": files_meta
    })


@app.post("/api/scan-upload")
def api_scan_upload():
    """
    Sube múltiples .docx (archivos o carpeta via webkitdirectory),
    calcula coincidencias contra 'whitelist' y retorna un token temporal.
    FormData:
      files=... (múltiples)
      whitelist[]=... (múltiples)
    """
    whitelist = request.form.getlist("whitelist[]") or []
    files = request.files.getlist("files")
    if not files:
        return jsonify({"ok": False, "error": "No enviaste .docx"}), 400

    token = uuid.uuid4().hex
    tmp_dir = os.path.join(TMP_ROOT, token)
    os.makedirs(tmp_dir, exist_ok=True)

    archivos = []
    for f in files:
        if not f.filename.lower().endswith(".docx"):
            continue
        raw = f.read()
        with open(os.path.join(tmp_dir, f.filename), "wb") as out:
            out.write(raw)
        archivos.append({"name": f.filename, "content": raw})

    files_meta = _collect_overview_whitelist(archivos, whitelist)
    return jsonify({
        "ok": True,
        "token": token,
        "count": len(archivos),
        "files_meta": files_meta
    })


@app.post("/api/merge")
def api_merge():
    """
    Unifica según modo y fuente.
    Body JSON:
      {
        "source": "folder" | "upload" | "upload_dir",
        "mode":   "unificado" | "grouped",
        "titles": ["t1","t2", ...],

        // si source = folder
        "folder": "/ruta/servidor",
        "include_subs": true/false,

        // si source = upload / upload_dir
        "token": "XXXX"
      }
    Devuelve: ZIP (unificado.zip o por_titulo.zip)
    """
    data = request.get_json(force=True)
    source = data.get("source")
    mode = (data.get("mode") or "unificado").lower()
    titles = data.get("titles") or []

    # Cargar archivos según la fuente
    archivos = []
    if source == "folder":
        folder = (data.get("folder") or "").strip()
        include_subs = bool(data.get("include_subs", True))
        if not folder or not os.path.isdir(folder):
            return jsonify({"ok": False, "error": "Carpeta inválida."}), 400
        archivos = _read_docx_from_folder(folder, include_subs)
    else:
        token = data.get("token")
        if not token:
            return jsonify({"ok": False, "error": "Falta token de subida."}), 400
        tmp_dir = os.path.join(TMP_ROOT, token)
        if not os.path.isdir(tmp_dir):
            return jsonify({"ok": False, "error": "Token no válido o expirado."}), 400
        for fn in os.listdir(tmp_dir):
            if fn.lower().endswith(".docx"):
                p = os.path.join(tmp_dir, fn)
                with open(p, "rb") as f:
                    archivos.append({"name": fn, "content": f.read()})

    if not archivos:
        return jsonify({"ok": False, "error": "No hay .docx para procesar."}), 400

    # Construir ZIP en memoria
    memzip = io.BytesIO()
    with zipfile.ZipFile(memzip, "w", zipfile.ZIP_DEFLATED) as zf:
        if mode == "unificado":
            # Unificado: incluimos secciones de todos los niveles, filtradas por títulos si se pasan
            res = procesar(
                archivos=archivos,
                niveles=list(range(1, 10)),
                titulos_whitelist=titles,
                enforce_whitelist=bool(titles)
            )
            for fname, data_bytes in res.items():
                zf.writestr(fname, data_bytes)
        else:
            # Por título: un .docx por cada título exacto (en cualquier nivel)
            total = 0
            for t in titles:
                partial = _merge_grouped_by_title_any_level(archivos, t)
                for fname, data_bytes in partial.items():
                    zf.writestr(fname, data_bytes)
                    total += 1
            if total == 0:
                return jsonify({"ok": False, "error": "No hubo coincidencias para los títulos seleccionados."}), 400

    memzip.seek(0)
    return send_file(
        memzip,
        mimetype="application/zip",
        as_attachment=True,
        download_name=("unificado.zip" if mode == "unificado" else "por_titulo.zip")
    )


@app.post("/api/cleanup")
def api_cleanup():
    """
    Limpia subidas temporales por token (ahorra disco en servidor).
    Body JSON: {"token": "XXXX"}
    """
    token = (request.get_json(force=True) or {}).get("token")
    if not token:
        return jsonify({"ok": True})
    tmp_dir = os.path.join(TMP_ROOT, token)
    if os.path.isdir(tmp_dir):
        shutil.rmtree(tmp_dir, ignore_errors=True)
    return jsonify({"ok": True})


if __name__ == "__main__":
    # Local
    app.run(host="0.0.0.0", port=8000, debug=False)
